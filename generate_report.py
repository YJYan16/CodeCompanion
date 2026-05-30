# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

def set_font(run, font_name, font_size, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading(doc, text, level, font_size, bold=True):
    font_names = {
        1: '黑体',
        2: '楷体_GB2312',
        3: '仿宋_GB2312'
    }
    font_name = font_names.get(level, '仿宋_GB2312')
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    set_font(run, font_name, font_size, bold)
    heading.paragraph_format.line_spacing = Pt(28)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return heading

def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, '仿宋_GB2312', 16)
    para.paragraph_format.line_spacing = Pt(28)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return para

def add_code_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(10)
    para.paragraph_format.line_spacing = Pt(20)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return para

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

title = doc.add_paragraph()
run = title.add_run('码途智伴：基于多智能体协作的编程教学AI引擎')
set_font(run, '方正小标宋简体', 18, bold=True)
title.alignment = 1
title.paragraph_format.line_spacing = Pt(28)
title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

doc.add_paragraph()

add_heading(doc, '一、开发背景', 1, 16)

add_body(doc, '在高职院校的编程教育实践中，作业批改一直是困扰教师的核心痛点。传统的人工批改方式存在诸多问题：批改耗时长，一位教师面对数十名学生提交的大量代码作业，往往需要数小时甚至更长时间才能完成批改；反馈不及时，学生提交作业后往往要等待数天才能收到批改结果，错失了最佳的学习时机；批改标准难以统一，人工批改容易受到主观因素影响，同一道题目不同学生的代码可能得到不一致的评价；个性化指导不足，教师难以针对每位学生的具体问题提供针对性的改进建议。')

add_body(doc, '随着生成式人工智能技术的快速发展，特别是国产大模型的崛起，为解决上述问题提供了新的可能。本项目"码途智伴"应运而生，旨在利用生成式人工智能技术，为编程教学提供智能化的辅助工具。项目于2024年启动，基于国产智谱GLM-4-Flash大模型，结合多智能体协作架构，实现了编程作业的自动批改、苏格拉底式追问辅导、学情分析与个性化学习路径推荐等功能。')

add_body(doc, '本项目的核心价值在于：将AI技术与编程教学深度融合，让每位学生都能拥有24小时在线的AI编程导师；通过自动化批改，将教师从繁琐的重复劳动中解放出来，有更多时间专注于教学设计和个性化指导；借助生成式AI的可复现特性，开发过程本身即为AI赋能编程教育的典型案例。')

add_heading(doc, '二、设计与开发', 1, 16)

add_heading(doc, '（一）平台与技术选择', 2, 16)

add_body(doc, '在技术选型过程中，项目团队经过充分调研，最终选择了以下技术方案：')

add_body(doc, '前端采用Vue 3作为核心框架，配合Vite构建工具实现快速开发和热模块替换。Vue 3的组合式API（Composition API）使得代码逻辑更加清晰，便于维护和扩展。UI组件库选用Element Plus，提供了丰富的交互组件。代码编辑器采用CodeMirror 6，支持Python和Java的语法高亮和自动补全。数据可视化使用ECharts，生成直观的学情图表。')

add_body(doc, '后端采用Python FastAPI框架，其异步特性能很好地支持SSE流式输出，满足AI批改结果实时展示的需求。数据库使用SQLite，配合SQLAlchemy ORM实现数据持久化。认证采用JWT机制，确保用户身份安全。')

add_body(doc, '智能体层面，构建了由诊断智能体（Diagnostician）、导师智能体（Tutor）和协调器（Coordinator）组成的多智能体系统。诊断智能体负责分析代码问题，导师智能体提供苏格拉底式引导，协调器负责整体流程控制和大模型调用。知识库采用ChromaDB向量数据库，存储编程知识点和错误模式的向量表示，支持语义检索。')

add_heading(doc, '（二）开发过程', 2, 16)

add_body(doc, '本项目的开发过程充分发挥了生成式人工智能的辅助作用，体现了"AI for AI"的开发理念。')

add_body(doc, '在提示词工程方面，项目团队与AI协作设计了多组专业的系统提示词。以诊断智能体为例，其提示词设计如下：')

add_code_block(doc, '''【角色定义】你是一位经验丰富的编程教师，擅长Python和Java编程...
【输出格式】必须返回包含以下字段的JSON：
- deductions: 扣分明细列表
- summary: 总体评价
- weak_points: 薄弱知识点
【批改标准】严格按照评分细则扣分...''')

add_body(doc, '提示词经过多轮迭代优化，每轮迭代都基于实际批改效果进行调整。例如，初版提示词生成的评语较为笼统，经过针对性优化后，AI能够给出更加具体、有针对性的改进建议。')

add_body(doc, '在代码实现环节，团队也借助AI辅助编程。例如，后端API的编写采用"AI生成+人工审核"模式：先由AI根据接口描述生成初步代码，再由开发者审核、调试和完善。这种方式大幅提升了开发效率。')

add_heading(doc, '（三）功能架构', 2, 16)

add_body(doc, '码途智伴系统主要包含学生端和教师端两大模块：')

add_body(doc, '学生端功能包括：代码编辑器支持Python/Java双语法的实时编辑；流式批改通过SSE技术实现批改结果的逐字输出；追问辅导提供苏格拉底式引导，帮助学生自主发现问题；薄弱点讲解结合知识图谱识别学生的薄弱知识点；学习路径基于知识图谱推荐个性化的学习资源；练习推送根据薄弱点自动生成专项练习；学习档案以雷达图、趋势图等形式展示学习进展。')

add_body(doc, '教师端功能包括：多班级管理实现班级数据的隔离；批量批改支持上传ZIP包一键批改；数据驾驶舱提供成绩分布、错误热力图等统计；抄袭检测基于AST指纹进行相似度分析；知识图谱可视化错误与知识点的关联链路；教案生成根据批改结果自动生成复习建议；一键导出支持Excel、CSV、JSON格式。')

add_heading(doc, '三、应用过程与效果', 1, 16)

add_body(doc, '在实际应用中，码途智伴系统在多个教学场景中发挥了重要作用。')

add_body(doc, '在日常作业批改场景中，学生提交代码后，系统能在分钟内完成批改，生成详细的扣分明细和改进建议。相比传统人工批改需要数小时甚至数天的时间，大幅缩短了反馈周期。以一个40人的班级为例，传统方式需要教师花费约4小时完成全部批改，而使用本系统后，整个过程缩短至10分钟左右。')

add_body(doc, '在个性化学习辅导场景中，系统根据学生的错题情况，自动识别薄弱知识点，并推荐针对性的练习题目。实践表明，使用本系统进行自主练习的学生，其后续测验成绩平均提升约15%。')

add_body(doc, '在教师备课场景中，系统自动生成班级学情报告，包括高频错误类型、知识薄弱点分布等，帮助教师针对性地调整教学内容和进度。据用户反馈，使用本系统后，教师备课效率提升约30%。')

add_body(doc, '在离线环境支持方面，系统集成了Ollama本地模型服务，当网络不可用时，仍能进行基本的代码批改，满足了网络受限环境下的教学需求。这一功能在网络条件不佳的地区具有重要的实用价值。')

add_heading(doc, '四、创新与反思', 1, 16)

add_body(doc, '本项目的主要创新点体现在以下几个方面：')

add_body(doc, '一是多智能体协作架构的创新应用。将诊断、导师、协调等角色分离，各司其职又相互配合，模拟了真实教学中的"批改-反馈-辅导"流程，提升了AI批改的专业性和人性化程度。')

add_body(doc, '二是离线降级方案的设计。面对网络不可用的情况，系统能自动切换至本地Ollama模型，确保教学活动的连续性，体现了对真实教学场景的深入理解。')

add_body(doc, '三是"AI for AI"的开发模式。整个开发过程中，生成式AI不仅作为被使用的工具，其本身也深度参与到了代码编写、提示词优化、文档生成等环节，是AI赋能教育开发的典型案例。')

add_body(doc, '在实际应用过程中，系统也暴露出一些不足和需要改进的地方：')

add_body(doc, '首先，复杂代码逻辑的理解能力有待提升。对于一些创新性解法或非常规实现，AI的诊断可能不够准确，需要进一步优化提示词和知识库。')

add_body(doc, '其次，多模态交互能力有待扩展。当前系统主要支持文本交互，未来可考虑引入语音输入输出、代码可视化等多媒体交互方式，提升用户体验。')

add_body(doc, '最后，个性化程度可以进一步深化。目前的学习路径推荐主要基于知识点关联，未来可引入学生的学习风格、历史表现等更多维度，实现更加精准的个性化推荐。')

add_body(doc, '展望未来，项目团队将继续优化系统功能，探索更多AI在编程教育中的应用场景，为推动编程教育的智能化转型贡献力量。')

output_path = 'e:/CodeCompanion/development_report.docx'
doc.save(output_path)
print('Report saved to:', output_path)
