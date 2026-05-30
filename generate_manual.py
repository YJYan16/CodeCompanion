# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

def set_font(run, font_name, font_size, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading(doc, text, level, font_size, bold=True):
    font_names = {
        1: '黑体',
        2: '楷体_GB2312',
        3: '仿宋_GB2312'
    }
    font_name = font_names.get(level, '仿宋_GB2312')
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    set_font(run, font_name, font_size, bold)
    heading.paragraph_format.line_spacing = Pt(28)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return heading

def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, '仿宋_GB2312', 16)
    para.paragraph_format.line_spacing = Pt(28)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return para

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

title = doc.add_paragraph()
run = title.add_run('码途智伴应用手册')
set_font(run, '方正小标宋简体', 18, bold=True)
title.alignment = 1
title.paragraph_format.line_spacing = Pt(28)
title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

doc.add_paragraph()

add_heading(doc, '一、系统简介', 1, 16)

add_body(doc, '码途智伴是一款面向高职院校编程教育的AI教学平台，基于多智能体协作架构实现编程作业的智能批改与个性化辅导。本系统融合了诊断智能体、导师智能体和协调器三大核心组件，能够自动分析学生代码、生成针对性评语、识别薄弱知识点并推荐学习路径。')

add_body(doc, '系统支持云端和本地双模型部署模式。在网络畅通时，使用智谱GLM-4-Flash大模型进行智能批改；在网络不可用时，可自动切换至Ollama本地部署的Qwen2.5-7B量化模型，确保教学活动的连续性。')

add_heading(doc, '二、学生端使用指南', 1, 16)

add_heading(doc, '（一）登录与界面', 2, 16)

add_body(doc, '学生使用个人账号登录系统，默认账号为学号，密码为123456。登录成功后进入学生端主界面，主要包含以下功能区域：左侧为题目列表，右侧为代码编辑器，下方为批改结果展示区。')

add_heading(doc, '（二）代码编辑与提交', 2, 16)

add_body(doc, '在代码编辑器中，学生可选择Python或Java语言编写代码。编辑器支持语法高亮和自动缩进。完成代码编写后，点击"提交批改"按钮，系统将实时展示批改结果。批改过程采用流式输出，学生可以直观看到AI逐字生成的评语。')

add_body(doc, '批改结果包含扣分明细、总体评价和薄弱知识点识别三部分。系统会指出代码中的具体问题，并给出改进建议。')

add_heading(doc, '（三）追问辅导', 2, 16)

add_body(doc, '学生对批改结果有疑问时，可点击"追问"按钮进入苏格拉底式辅导模式。系统通过启发式提问引导学生自主发现问题所在，而非直接给出答案。这种教学方式有助于培养学生的独立思考能力和问题解决能力。')

add_heading(doc, '（四）练习模式', 2, 16)

add_body(doc, '系统根据学生的薄弱知识点自动生成专项练习。学生在练习模式下完成练习题后，系统会自动批改并记录学习数据。练习完成后可切换回普通作业模式继续正常学习。切换时系统会保留原代码状态。')

add_heading(doc, '（五）学习档案', 2, 16)

add_body(doc, '学习档案页面以雷达图、趋势图等形式展示学生的学习进展。学生可以查看各知识点的掌握情况、历史成绩变化趋势以及教师的个性化建议。')

add_heading(doc, '（六）离线模式', 2, 16)

add_body(doc, '当网络连接不可用时，系统界面右上角会显示"离线模式"提示。此时系统自动切换至本地Ollama模型进行代码批改。虽然本地模型的响应可能稍慢，但基本批改功能仍可正常使用。重新联网后，系统会自动恢复云端模型。')

add_heading(doc, '三、教师端使用指南', 1, 16)

add_heading(doc, '（一）班级管理', 2, 16)

add_body(doc, '教师登录后进入教师端，可创建和管理多个班级。每个班级的学生数据相互隔离，教师可随时切换当前管理的班级。')

add_heading(doc, '（二）作业发布', 2, 16)

add_body(doc, '教师可在题目管理页面创建新题目，设置题目描述、参考答案和评分细则。题目创建后可分配给指定班级，学生即可在学生端看到并完成作业。')

add_heading(doc, '（三）批量批改', 2, 16)

add_body(doc, '对于需要人工批改的复杂作业，教师可上传学生提交的代码压缩包（ZIP格式），系统会自动解压并逐一进行AI预批改，生成批改建议供教师参考，大幅提升批改效率。')

add_heading(doc, '（四）数据驾驶舱', 2, 16)

add_body(doc, '数据驾驶舱提供班级学情的全景视图，包括：成绩分布直方图、高频错误类型统计、知识点薄弱点热力图等。教师可据此调整教学计划和重点。')

add_heading(doc, '（五）一键导出', 2, 16)

add_body(doc, '教师可批量导出班级成绩数据，支持Excel、CSV、JSON等多种格式。同时支持导出学生代码包，便于存档和复审。')

add_heading(doc, '四、常见问题', 1, 16)

add_body(doc, 'Q1：批改结果不符合预期怎么办？')
add_body(doc, 'A1：批改结果仅供参考，教师可在教师端对自动批改结果进行调整。如发现AI理解有误，可通过追问功能进一步澄清，或反馈给开发团队优化提示词。')

add_body(doc, 'Q2：离线模式功能受限？')
add_body(doc, 'A2：离线模式下系统使用本地模型进行批改，功能相对基础但可满足基本批改需求。建议在网络正常时使用云端模型以获得更精准的批改效果。')

add_body(doc, 'Q3：如何获取更详细的学情报告？')
add_body(doc, 'A3：在教师端数据驾驶舱页面，可按班级、按题目等多种维度查看学情分析报告，支持时间范围筛选和数据导出。')

output_path = 'e:/CodeCompanion/应用手册.docx'
doc.save(output_path)
print('Manual saved to:', output_path)